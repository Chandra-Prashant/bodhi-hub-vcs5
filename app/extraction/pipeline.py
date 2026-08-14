"""
Extraction pipeline — Phase 3.

Reads a project document and produces a validated `ExtractionResult`. The model
call is isolated behind `Extractor.complete()` so the parsing, validation and
confidence logic — which is where the behaviour that matters lives — is
testable without a model or an API key.

The model's only job here is locating values in text. It is given a fixed JSON
schema and told, explicitly, not to compute anything. Everything downstream of
the response is deterministic Python.
"""

from __future__ import annotations

import json
from abc import ABC, abstractmethod
from datetime import date
from pathlib import Path

from app.extraction.guards import assert_extraction_safe
from app.extraction.schema import (
    Confidence,
    ExtractedField,
    ExtractionResult,
    ExtractionStatus,
    ProjectExtraction,
    band,
)

# Checked at import: a schema change that crosses the calculated-field boundary
# fails here rather than on a client document.
assert_extraction_safe(ProjectExtraction.model_fields)


SYSTEM_PROMPT = """\
You extract stated values from project documents. You do not calculate.

Return JSON matching the given schema exactly. For each field return:
  value        the value exactly as the document states it, or null
  score        0.0-1.0, how certain you are this is the right value
  source_page  the page number you found it on, or null
  source_text  the sentence or table row you took it from, verbatim
  note         anything a human reviewer should know

Rules, without exception:
- Never compute, derive, convert, sum, or estimate a value. If a figure is not
  stated in the document, return null. A field you had to work out is a field
  you must leave empty.
- Never convert units. Return the number as written and note the unit in
  `note`. Unit conversion is a calculation and happens downstream.
- Never repair a value that looks wrong. Return it as stated with a low score
  and say why in `note`. Correcting it hides an error a reviewer needs to see.
- If the document states a figure that is clearly a calculated result — total
  emission reductions, an IRR, a grid emission factor — do not return it. Those
  are computed by the system, not read from documents.
- `source_text` must be copied verbatim from the document. A reviewer uses it
  to check you without reopening the file.
- Lower your score when the value is ambiguous, appears more than once with
  different values, or you inferred it from context rather than reading it.
"""


class Extractor(ABC):
    """The model boundary. Everything else in this module is deterministic."""

    name = "abstract"

    @abstractmethod
    def complete(self, system: str, document_text: str) -> str:
        """Return the model's raw JSON response."""


class GeminiExtractor(Extractor):
    name = "gemini"

    def __init__(self, api_key: str, model: str = "gemini-flash-latest") -> None:
        self.api_key = api_key
        self.model = model
        self.name = model

    def complete(self, system: str, document_text: str) -> str:
        from google import genai

        client = genai.Client(api_key=self.api_key)
        response = client.models.generate_content(
            model=self.model,
            contents=f"{system}\n\n--- DOCUMENT ---\n{document_text}",
            config={"response_mime_type": "application/json"},
        )
        return response.text


# ---------------------------------------------------------------------------
# Document loading
# ---------------------------------------------------------------------------


class UnsupportedDocument(Exception):
    pass


def load_text(path: Path) -> tuple[str, int]:
    """Return (text, page_count).

    Raises rather than returning empty text for a readable-but-empty document:
    a scanned PDF with no text layer needs OCR, and silently treating it as an
    empty document would produce a confident extraction of nothing.
    """
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "") for page in reader.pages]
        text = "\n\n".join(
            f"[page {i}]\n{content}" for i, content in enumerate(pages, 1)
        )
        if not any(p.strip() for p in pages):
            raise UnsupportedDocument(
                f"{path.name} has no extractable text layer. It is probably a "
                f"scan and needs OCR before extraction.")
        return text, len(pages)

    if suffix in {".docx", ".doc"}:
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        parts += [
            " | ".join(c.text.strip() for c in row.cells)
            for table in document.tables for row in table.rows
        ]
        if not parts:
            raise UnsupportedDocument(f"{path.name} contains no text.")
        return "\n".join(parts), 1

    if suffix in {".txt", ".md"}:
        text = path.read_text(errors="replace")
        if not text.strip():
            raise UnsupportedDocument(f"{path.name} is empty.")
        return text, 1

    raise UnsupportedDocument(
        f"{path.suffix or 'file'} is not a supported document type. "
        f"Supported: .pdf, .docx, .txt, .md")


# ---------------------------------------------------------------------------
# Response handling
# ---------------------------------------------------------------------------


def _coerce_date(raw: object) -> date | None:
    if isinstance(raw, date):
        return raw
    if isinstance(raw, str) and raw.strip():
        for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y", "%d-%b-%Y", "%d %B %Y"):
            try:
                from datetime import datetime

                return datetime.strptime(raw.strip(), fmt).date()
            except ValueError:
                continue
    return None


def parse_response(payload: str) -> ProjectExtraction:
    """Turn a model response into a validated ProjectExtraction.

    Unknown keys are dropped rather than raising: a model returning an extra
    field should not fail a document, and the schema forbids extras precisely
    so they cannot slip through unnoticed here.
    """
    text = payload.strip()
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]

    raw = json.loads(text)
    if not isinstance(raw, dict):
        raise ValueError("Extraction response was not a JSON object.")

    fields: dict[str, ExtractedField] = {}
    for name, annotation in ProjectExtraction.model_fields.items():
        entry = raw.get(name)
        if not isinstance(entry, dict):
            fields[name] = ExtractedField()
            continue

        value = entry.get("value")
        if name == "initial_crediting_period_start":
            value = _coerce_date(value)
        elif value is not None:
            value = str(value).strip() or None

        try:
            score = float(entry.get("score", 0.0))
        except (TypeError, ValueError):
            score = 0.0
        score = min(max(score, 0.0), 1.0)

        page = entry.get("source_page")
        fields[name] = ExtractedField(
            value=value,
            confidence=band(score) if value is not None else Confidence.NOT_FOUND,
            score=score if value is not None else 0.0,
            source_page=int(page) if isinstance(page, (int, float)) else None,
            source_text=str(entry.get("source_text", ""))[:600],
            note=str(entry.get("note", ""))[:400],
        )

    return ProjectExtraction(**fields)


def extract(
    path: Path,
    extractor: Extractor,
    minimum_fields: int = 3,
) -> ExtractionResult:
    """Run the pipeline over one document.

    Never raises for an expected failure. An unreadable document, an
    unparseable response or a near-empty extraction all come back as a FAILED
    result carrying its reason, because Rules.md requires the document to be
    flagged for manual entry rather than the failure disappearing.
    """
    try:
        text, pages = load_text(path)
    except UnsupportedDocument as exc:
        return ExtractionResult(
            document_name=path.name, status=ExtractionStatus.FAILED,
            error=str(exc), model=extractor.name)
    except Exception as exc:  # noqa: BLE001 — any reader failure flags the doc
        return ExtractionResult(
            document_name=path.name, status=ExtractionStatus.FAILED,
            error=f"Could not read {path.name}: {exc}", model=extractor.name)

    try:
        response = extractor.complete(SYSTEM_PROMPT, text)
    except Exception as exc:  # noqa: BLE001
        return ExtractionResult(
            document_name=path.name, status=ExtractionStatus.FAILED,
            error=f"Extraction model failed: {exc}",
            pages_read=pages, model=extractor.name)

    try:
        data = parse_response(response)
    except Exception as exc:  # noqa: BLE001
        return ExtractionResult(
            document_name=path.name, status=ExtractionStatus.FAILED,
            error=f"Extraction response could not be parsed: {exc}",
            pages_read=pages, model=extractor.name)

    found = data.fields_found()
    if len(found) < minimum_fields:
        return ExtractionResult(
            document_name=path.name, status=ExtractionStatus.FAILED,
            data=data, pages_read=pages, model=extractor.name,
            error=(
                f"Only {len(found)} field(s) extracted, below the minimum of "
                f"{minimum_fields}. This document probably is not a project "
                f"description, or needs manual entry."))

    clean = not data.fields_needing_review() and not data.missing_required()
    status = ExtractionStatus.EXTRACTED if clean else ExtractionStatus.PARTIAL
    return ExtractionResult(
        document_name=path.name, status=status, data=data,
        pages_read=pages, model=extractor.name)
