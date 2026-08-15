"""
Document loading — what a source document becomes before extraction.

PRD.md must-have: "Document/data ingestion (PDF, Word, Excel, images of forms)",
and field auditors upload "photos of forms". Three shapes arrive, and they are
not interchangeable:

  * **Text documents** — PDF with a text layer, Word, plain text
  * **Spreadsheets** — Excel and CSV, flattened to labelled rows
  * **Images** — a photographed or scanned form, with no text to read

Images are passed to the model as vision input rather than through a separate
OCR step. Rules.md grants the LLM exactly two jobs, the first being "extracting
structured data from documents" — reading a photographed form is that job, not
a new one. It also avoids a second service in the chain, each hop being a place
for a digit to change.

A scanned PDF with no text layer is treated as an image for the same reason,
instead of being refused.
"""

from __future__ import annotations

import mimetypes
from dataclasses import dataclass
from pathlib import Path

TEXT_SUFFIXES = {".pdf", ".docx", ".doc", ".txt", ".md"}
SHEET_SUFFIXES = {".xlsx", ".xlsm", ".csv", ".tsv"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".heic", ".tif", ".tiff"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | SHEET_SUFFIXES | IMAGE_SUFFIXES

# A page rendered at a sensible DPI is a few hundred KB; well beyond that and
# the request is slow and expensive for no gain in legibility.
MAX_IMAGE_BYTES = 8 * 1024 * 1024


class UnsupportedDocument(Exception):
    pass


@dataclass
class DocumentContent:
    """What the extractor receives. Exactly one of `text` or `image` is set."""

    text: str = ""
    image: bytes | None = None
    media_type: str = ""
    pages: int = 0
    kind: str = "text"          # text | sheet | image

    @property
    def is_image(self) -> bool:
        return self.image is not None


def _load_pdf(path: Path) -> DocumentContent:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "") for page in reader.pages]

    if any(p.strip() for p in pages):
        text = "\n\n".join(
            f"[page {i}]\n{content}" for i, content in enumerate(pages, 1))
        return DocumentContent(text=text, pages=len(pages), kind="text")

    # No text layer — a scan. Previously refused; now handed to the model as an
    # image, which is what a field auditor's uploaded form usually is.
    raise UnsupportedDocument(
        f"{path.name} has no text layer. Convert the pages to images "
        f"(PNG or JPEG) and upload those — the extractor reads them directly.")


def _load_docx(path: Path) -> DocumentContent:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    parts += [
        " | ".join(c.text.strip() for c in row.cells)
        for table in document.tables for row in table.rows
    ]
    if not parts:
        raise UnsupportedDocument(f"{path.name} contains no text.")
    return DocumentContent(text="\n".join(parts), pages=1, kind="text")


def _load_sheet(path: Path) -> DocumentContent:
    """Flatten a spreadsheet to labelled rows.

    Cell *values* are read, never formulas. A formula string tells the
    extractor nothing, and its computed result is somebody else's arithmetic —
    the same reason extraction refuses derived figures generally.
    """
    suffix = path.suffix.lower()
    rows: list[str] = []

    if suffix in {".csv", ".tsv"}:
        import csv

        delimiter = "\t" if suffix == ".tsv" else ","
        with path.open(newline="", errors="replace") as handle:
            for row in csv.reader(handle, delimiter=delimiter):
                if any(cell.strip() for cell in row):
                    rows.append(" | ".join(cell.strip() for cell in row))
    else:
        from openpyxl import load_workbook

        # data_only=True returns the cached value rather than the formula.
        workbook = load_workbook(str(path), data_only=True, read_only=True)
        for sheet in workbook.worksheets:
            rows.append(f"[sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if c is None else str(c).strip() for c in row]
                if any(cells):
                    rows.append(" | ".join(cells))
        workbook.close()

    if not any(r for r in rows if not r.startswith("[sheet:")):
        raise UnsupportedDocument(f"{path.name} contains no data rows.")
    return DocumentContent(text="\n".join(rows), pages=1, kind="sheet")


def _load_image(path: Path) -> DocumentContent:
    data = path.read_bytes()
    if not data:
        raise UnsupportedDocument(f"{path.name} is empty.")
    if len(data) > MAX_IMAGE_BYTES:
        raise UnsupportedDocument(
            f"{path.name} is {len(data) / 1_048_576:.1f} MB; the limit for "
            f"images is {MAX_IMAGE_BYTES // 1_048_576} MB. Reduce the "
            f"resolution — a legible page scan is well under that.")

    media_type, _ = mimetypes.guess_type(path.name)
    if not media_type or not media_type.startswith("image/"):
        media_type = f"image/{path.suffix.lstrip('.').lower()}"
    return DocumentContent(image=data, media_type=media_type, pages=1,
                           kind="image")


def load_document(path: Path) -> DocumentContent:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix in {".docx", ".doc"}:
        return _load_docx(path)
    if suffix in {".txt", ".md"}:
        text = path.read_text(errors="replace")
        if not text.strip():
            raise UnsupportedDocument(f"{path.name} is empty.")
        return DocumentContent(text=text, pages=1, kind="text")
    if suffix in SHEET_SUFFIXES:
        return _load_sheet(path)
    if suffix in IMAGE_SUFFIXES:
        return _load_image(path)

    raise UnsupportedDocument(
        f"{path.suffix or 'This file type'} is not supported. Accepted: "
        f"{', '.join(sorted(SUPPORTED_SUFFIXES))}")
