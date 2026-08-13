"""
Module 3b — docx mechanics.

Fills the official Verra template **as shipped**. No manual tagging step and no
`{{ }}` placeholders, which matters because Verra reissues these templates: a
new version drops into app/templates/ and keeps working, whereas a hand-tagged
copy would have to be re-tagged every time.

Two structural facts about the template make this possible:

  1. Structured data lives in two-column tables whose left cell is the field
     label ("Project name", "Crediting period start", ...).
  2. Every guidance block uses the paragraph style `Instruction` — 232 of them
     in the v5.0A Project Description. Drafted prose replaces the Instruction
     paragraphs under its heading; whatever remains is unfilled and is reported.

That last point is the useful one: leftover `Instruction` paragraphs are a
precise, mechanical inventory of what a human still has to write. It is the
seed of the auditor agent in Module 7.
"""

from __future__ import annotations

import copy
from dataclasses import dataclass

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

INSTRUCTION_STYLE = "Instruction"
HEADING_STYLES = ("Heading 1", "Heading 2", "Heading 3", "Heading 4")


@dataclass
class FillReport:
    fields_written: list[str]
    fields_not_found: list[str]
    sections_written: list[str]
    sections_not_found: list[str]
    instructions_remaining: dict[str, int]

    @property
    def completion_ratio(self) -> float:
        total = len(self.fields_written) + len(self.fields_not_found)
        return len(self.fields_written) / total if total else 0.0


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace a paragraph's text while keeping its first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _set_cell_value(cell, text: str) -> None:
    """Write into a table cell, preserving the cell's paragraph formatting."""
    target = cell.paragraphs[0]
    _set_paragraph_text(target, text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _insert_paragraph_after(paragraph: Paragraph, text: str,
                            style: str | None = None) -> Paragraph:
    new_el = copy.deepcopy(paragraph._element)
    paragraph._element.addnext(new_el)
    new_par = Paragraph(new_el, paragraph._parent)
    for run in new_par.runs[1:]:
        run._element.getparent().remove(run._element)
    _set_paragraph_text(new_par, text)
    if style:
        new_par.style = style
    return new_par


def _is_heading(paragraph: Paragraph) -> bool:
    return paragraph.style.name in HEADING_STYLES


# ---------------------------------------------------------------------------
# Field filling
# ---------------------------------------------------------------------------

def fill_fields(doc: Document, fields: dict[str, str]) -> tuple[list[str], list[str]]:
    """Match the left cell of every two-column row against the field labels.

    Matching is case-insensitive on the stripped label and takes the FIRST
    match only — several labels ("Justification") repeat throughout the
    template, and writing the same value into all of them would be worse than
    leaving them blank.
    """
    written: list[str] = []
    remaining = dict(fields)

    def walk(tables: list[Table]) -> None:
        for table in tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue
                label = cells[0].text.strip().rstrip(":")
                for key in list(remaining):
                    if label.lower() == key.lower():
                        _set_cell_value(cells[1], remaining.pop(key))
                        written.append(key)
                        break
                for cell in cells:
                    if cell.tables:
                        walk(cell.tables)

    walk(doc.tables)
    return written, list(remaining)


# ---------------------------------------------------------------------------
# Section filling
# ---------------------------------------------------------------------------

def _section_paragraphs(doc: Document, heading_text: str) -> tuple[Paragraph | None, list[Paragraph]]:
    """Return the heading paragraph and the Instruction paragraphs beneath it."""
    paragraphs = doc.paragraphs
    heading = None
    start = -1
    for i, p in enumerate(paragraphs):
        if _is_heading(p) and p.text.strip().lower() == heading_text.strip().lower():
            heading, start = p, i
            break
    if heading is None:
        return None, []

    body: list[Paragraph] = []
    for p in paragraphs[start + 1:]:
        if _is_heading(p):
            break
        if p.style.name == INSTRUCTION_STYLE:
            body.append(p)
    return heading, body


def fill_sections(
    doc: Document,
    sections: dict[str, list[str]],
) -> tuple[list[str], list[str]]:
    """Replace guidance under each heading with the drafted paragraphs."""
    written: list[str] = []
    not_found: list[str] = []

    for heading_text, content in sections.items():
        heading, instructions = _section_paragraphs(doc, heading_text)
        if heading is None:
            not_found.append(heading_text)
            continue
        if not content:
            continue

        if instructions:
            # Reuse the first guidance paragraph as the anchor, then chain.
            anchor = instructions[0]
            anchor.style = doc.styles["Normal"]
            _set_paragraph_text(anchor, content[0])
            for text in content[1:]:
                anchor = _insert_paragraph_after(anchor, text, "Normal")
            for leftover in instructions[1:]:
                _delete_paragraph(leftover)
        else:
            anchor = heading
            for text in content:
                anchor = _insert_paragraph_after(anchor, text, "Normal")

        written.append(heading_text)

    return written, not_found


# ---------------------------------------------------------------------------
# Estimated reductions table
# ---------------------------------------------------------------------------

def fill_estimates_table(doc: Document, estimates: list) -> bool:
    """Populate the summary table of estimated reductions by year.

    Located by header text rather than index: table ordering shifts between
    template revisions, and an index-based lookup would silently write into the
    wrong table after an update.
    """
    if not estimates:
        return False

    for table in doc.tables:
        if not table.rows:
            continue
        header = " ".join(c.text.strip().lower() for c in table.rows[0].cells)
        if "year" in header and (
                "estimated" in header or "reduction" in header or "removal" in header):
            template_row = table.rows[-1]
            for est in estimates:
                new_row = copy.deepcopy(template_row._tr)
                table._tbl.append(new_row)
                cells = table.rows[-1].cells
                values = [
                    est.period_label,
                    f"{est.baseline_tco2e:,.0f}",
                    f"{est.project_tco2e:,.0f}",
                    f"{est.leakage_tco2e:,.0f}",
                    f"{est.reductions_tco2e:,.0f}",
                ]
                for cell, value in zip(cells, values[:len(cells)]):
                    _set_cell_value(cell, value)
            _delete_row(table, template_row)
            return True
    return False


def _delete_row(table: Table, row) -> None:
    row._tr.getparent().remove(row._tr)


# ---------------------------------------------------------------------------
# Completion audit
# ---------------------------------------------------------------------------

def remaining_instructions(doc: Document) -> dict[str, int]:
    """Count unfilled guidance blocks per section heading.

    This is the completion report a project manager works from — and the input
    the Module 7 auditor agent will reason over.
    """
    counts: dict[str, int] = {}
    current = "(document preamble)"
    for p in doc.paragraphs:
        if _is_heading(p):
            current = p.text.strip()
        elif p.style.name == INSTRUCTION_STYLE and p.text.strip():
            counts[current] = counts.get(current, 0) + 1
    return counts


def strip_instructions(doc: Document) -> int:
    """Remove all remaining guidance text.

    Only for a final submission copy. Run it on a working draft and you destroy
    the record of what has not yet been written.
    """
    removed = 0
    for p in list(doc.paragraphs):
        if p.style.name == INSTRUCTION_STYLE:
            _delete_paragraph(p)
            removed += 1
    return removed


# ---------------------------------------------------------------------------
# Appendix 2 — Data and Parameters
# ---------------------------------------------------------------------------
#
# The template ships ONE blank Data/Parameter table for A2.1 and one for A2.2.
# Each parameter needs its own, so the blank is cloned per parameter and the
# original removed. Tables are located by their row-label signature rather than
# by index: index-based lookup breaks silently the first time Verra reissues
# the template with a section inserted above.

_AT_VALIDATION_KEYS = {"data/parameter name", "data/parameter"}
_MONITORED_MARKER = "frequency of monitoring/recording"

_ROW_MAP = {
    "data/parameter name": "name",
    "data/parameter": "name",
    "data unit": "unit",
    "description": "description",
    "equation number": "equations",
    "equation number#": "equations",
    "source of data": "source_of_data",
    "value applied": "value_applied",
    "justification of value applied": "justification",
    "description of measurement methods and procedures to be applied":
        "measurement_methods",
    "frequency of monitoring/recording": "monitoring_frequency",
    "monitoring equipment": "monitoring_equipment",
    "qa/qc procedures to be applied": "qa_qc",
    "purpose of data": "purpose",
    "calculation method": "calculation_method",
    "comments": "comments",
}


def _table_signature(table: Table) -> tuple[str, ...]:
    return tuple(r.cells[0].text.strip().lower()
                 for r in table.rows if len(r.cells) >= 2)


def _find_parameter_tables(doc: Document) -> tuple[Table | None, Table | None]:
    at_validation = monitored = None
    for table in doc.tables:
        sig = _table_signature(table)
        if not sig or sig[0] not in _AT_VALIDATION_KEYS:
            continue
        if _MONITORED_MARKER in sig:
            monitored = monitored or table
        else:
            at_validation = at_validation or table
    return at_validation, monitored


def _fill_parameter_table(table: Table, parameter) -> None:
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip().lower()
        attr = _ROW_MAP.get(label)
        if attr is None:
            continue
        value = getattr(parameter, attr, "")
        if value:
            _set_cell_value(row.cells[1], str(value))


def _clone_table_after(table: Table) -> Table:
    """Copy a table and insert it after the original, with a spacer paragraph.

    Word merges directly adjacent tables into one, so the empty paragraph
    between them is structural, not cosmetic.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    spacer = OxmlElement("w:p")
    table._tbl.addnext(spacer)
    new_tbl = copy.deepcopy(table._tbl)
    spacer.addnext(new_tbl)
    return Table(new_tbl, table._parent)


def fill_parameter_tables(
    doc: Document,
    at_validation: list,
    monitored: list,
) -> dict[str, int]:
    """Render one Data/Parameter table per parameter in Appendix 2."""
    written = {"at_validation": 0, "monitored": 0}
    av_template, mon_template = _find_parameter_tables(doc)

    for template, params, key in (
        (av_template, at_validation, "at_validation"),
        (mon_template, monitored, "monitored"),
    ):
        if template is None or not params:
            continue
        anchor = template
        for parameter in params:
            clone = _clone_table_after(anchor)
            _fill_parameter_table(clone, parameter)
            anchor = clone
            written[key] += 1
        template._tbl.getparent().remove(template._tbl)

    return written
