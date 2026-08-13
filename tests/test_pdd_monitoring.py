"""Module 4b tests — Appendix 2 rendering and the Monitoring Plan section."""

from __future__ import annotations

from datetime import date

import docx
import pytest

from app.domain import constants as K
from app.domain.classification import ProjectIntake, Severity, classify
from app.domain.monitoring import build_monitoring_parameters, monitoring_plan_sections
from app.domain.pdd_content import ProjectIdentity, build_pdd_content
from app.services import docx_filler
from app.services.pdd_builder import build_pdd


def _intake(**overrides) -> ProjectIntake:
    base = dict(
        name="Aligarh Solar One", proponent="Bodhi Hub Client",
        country_iso2="IN", technology=K.Technology.SOLAR_PV_TERRESTRIAL,
        installed_capacity_mw=50.0, expected_annual_generation_mwh=87_600.0,
        initial_crediting_period_start=date(2026, 3, 1),
    )
    base.update(overrides)
    return ProjectIntake(**base)


def _content(technology=K.Technology.SOLAR_PV_TERRESTRIAL, has_bess=False,
             country="IN", **kw):
    intake = _intake(technology=technology, country_iso2=country, **kw)
    mp = build_monitoring_parameters(technology, has_bess=has_bess)
    return build_pdd_content(intake, classify(intake), ProjectIdentity(),
                             monitoring=mp)


def _parameter_tables(path):
    doc = docx.Document(str(path))
    return [t for t in doc.tables
            if docx_filler._table_signature(t)
            and docx_filler._table_signature(t)[0]
            in ("data/parameter name", "data/parameter")]


def _finding(findings, check):
    return next(f for f in findings if f.check == check)


# --- Appendix 2 rendering --------------------------------------------------

def test_one_table_is_rendered_per_parameter(tmp_path):
    content = _content()
    build_pdd(content, tmp_path / "PDD.docx")
    tables = _parameter_tables(tmp_path / "PDD.docx")
    assert len(tables) == len(content.monitoring.all())


def test_parameter_names_land_in_the_tables(tmp_path):
    build_pdd(_content(), tmp_path / "PDD.docx")
    names = {t.rows[0].cells[1].text.strip()
             for t in _parameter_tables(tmp_path / "PDD.docx")}
    assert {"EFembodied", "EFgrid,CM,y", "EGfacility,y", "EGPJ_Add,y"} <= names


def test_the_blank_template_table_is_removed(tmp_path):
    """A leftover empty Data/Parameter table reads as an unfilled requirement."""
    build_pdd(_content(), tmp_path / "PDD.docx")
    for table in _parameter_tables(tmp_path / "PDD.docx"):
        assert table.rows[0].cells[1].text.strip() != ""


def test_mandated_embodied_default_appears_in_the_document(tmp_path):
    build_pdd(_content(), tmp_path / "PDD.docx")
    table = next(t for t in _parameter_tables(tmp_path / "PDD.docx")
                 if t.rows[0].cells[1].text.strip() == "EFembodied")
    values = " ".join(c.text for r in table.rows for c in r.cells)
    assert "43" in values


def test_wind_renders_its_own_embodied_default(tmp_path):
    build_pdd(_content(K.Technology.WIND_ONSHORE), tmp_path / "PDD.docx")
    table = next(t for t in _parameter_tables(tmp_path / "PDD.docx")
                 if t.rows[0].cells[1].text.strip() == "EFembodied")
    values = " ".join(c.text for r in table.rows for c in r.cells)
    assert "13" in values


def test_monitored_tables_carry_frequency_and_qa_qc(tmp_path):
    build_pdd(_content(), tmp_path / "PDD.docx")
    table = next(t for t in _parameter_tables(tmp_path / "PDD.docx")
                 if t.rows[0].cells[1].text.strip() == "EGfacility,y")
    labelled = {r.cells[0].text.strip().lower(): r.cells[1].text.strip()
                for r in table.rows if len(r.cells) >= 2}
    assert "monthly" in labelled["frequency of monitoring/recording"].lower()
    assert "calibrat" in labelled["qa/qc procedures to be applied"].lower()


def test_hydro_renders_the_reservoir_parameter(tmp_path):
    content = _content(K.Technology.HYDRO, country="NP",
                       installed_capacity_mw=12.0,
                       expected_annual_generation_mwh=40_000.0)
    build_pdd(content, tmp_path / "PDD.docx")
    names = {t.rows[0].cells[1].text.strip()
             for t in _parameter_tables(tmp_path / "PDD.docx")}
    assert "EFRes" in names


def test_bess_renders_fire_suppression_parameters(tmp_path):
    content = _content(has_bess=True)
    build_pdd(content, tmp_path / "PDD.docx")
    names = {t.rows[0].cells[1].text.strip()
             for t in _parameter_tables(tmp_path / "PDD.docx")}
    assert {"GWPagent", "Me,released,y"} <= names


def test_tables_are_separated_so_word_does_not_merge_them(tmp_path):
    """Adjacent tables with no paragraph between them merge into one on open."""
    build_pdd(_content(), tmp_path / "PDD.docx")
    doc = docx.Document(str(tmp_path / "PDD.docx"))
    body = list(doc.element.body)
    tags = [e.tag.split("}")[-1] for e in body]
    for i in range(len(tags) - 1):
        assert not (tags[i] == "tbl" and tags[i + 1] == "tbl")


# --- monitoring plan prose -------------------------------------------------

def test_monitoring_plan_covers_the_required_ground():
    sections = monitoring_plan_sections(K.Technology.SOLAR_PV_TERRESTRIAL)
    text = " ".join(sections["Monitoring Plan"]).lower()
    for expected in ("egfacility", "vt0011", "tool03", "calibrat",
                     "interpolation", "monthly"):
        assert expected in text


def test_monitoring_plan_states_record_retention():
    text = " ".join(monitoring_plan_sections(
        K.Technology.WIND_ONSHORE)["Monitoring Plan"]).lower()
    assert "two years" in text


def test_bess_adds_the_fire_suppression_narrative():
    plain = monitoring_plan_sections(K.Technology.SOLAR_PV_TERRESTRIAL)
    with_bess = monitoring_plan_sections(
        K.Technology.SOLAR_PV_TERRESTRIAL, has_bess=True)
    assert len(with_bess["Monitoring Plan"]) > len(plain["Monitoring Plan"])
    assert "fire suppression" in " ".join(
        with_bess["Monitoring Plan"]).lower()


def test_monitoring_personnel_section_is_drafted():
    sections = monitoring_plan_sections(K.Technology.SOLAR_PV_TERRESTRIAL)
    assert sections["Monitoring Personnel"]


def test_monitoring_sections_reach_the_rendered_document(tmp_path):
    result = build_pdd(_content(), tmp_path / "PDD.docx")
    assert "Monitoring Plan" in result.report.sections_written
    assert "Monitoring Personnel" in result.report.sections_written


# --- refusal without monitoring data --------------------------------------

def test_missing_monitoring_parameters_block_the_pdd():
    intake = _intake()
    content = build_pdd_content(intake, classify(intake), ProjectIdentity())
    assert _finding(content.findings, "pdd.monitoring").severity is Severity.FAIL


def test_bess_warning_propagates_to_the_pdd():
    content = _content(has_bess=True)
    assert _finding(content.findings, "vmr0017.bess").severity is Severity.WARNING


def test_embodied_ef_finding_propagates_to_the_pdd():
    content = _content()
    f = _finding(content.findings, "vmr0017.embodied_ef")
    assert f.severity is Severity.PASS
    assert "43" in f.message
