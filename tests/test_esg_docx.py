"""
ESG docx rendering tests.

The dropdown tests carry the weight. Writing a value into the cell text instead
of into the content control leaves Word showing "Select" — the document looks
filled in a text dump and is blank to a reviewer opening it.
"""

from __future__ import annotations

import docx
import pytest
from docx.oxml.ns import qn

from app.domain.esg import RiskCategory, RiskEntry, assess_esg
from app.services.esg_docx import (
    LIKELIHOOD_OPTIONS,
    SEVERITY_OPTIONS,
    _sdt_elements,
    dropdown_options,
    find_category_tables,
    render_esg,
)


def _entry(category, **kw) -> RiskEntry:
    base = dict(
        category=category,
        risk_id=f"{category.value}.1",
        description="Temporary habitat disturbance during construction.",
        severity=2,
        likelihood=3,
        justification="Site is degraded agricultural land; survey completed.",
        mitigation="Works confined to the fenced footprint; ecologist on call.",
    )
    base.update(kw)
    return RiskEntry(**base)


def _assessment(**per_category):
    return assess_esg([
        _entry(c, **per_category.get(c, {})) for c in RiskCategory
    ])


@pytest.fixture
def rendered(tmp_path):
    path = tmp_path / "esg.docx"
    report = render_esg(_assessment(), path)
    return docx.Document(str(path)), report


# --- template structure ----------------------------------------------------

def test_category_tables_are_found_by_heading_code():
    """Matched by the '(S3)' in the heading rather than table index, so
    inserting a section upstream cannot silently shift the mapping."""
    from app.services.pdd_builder import _resolve

    doc = docx.Document(str(_resolve("VCS-ESG-Risk-Assessment-Template-v5.0.docx")))
    tables = find_category_tables(doc)
    assert RiskCategory.E1 in tables
    assert RiskCategory.S3 in tables
    assert RiskCategory.G2 in tables


def test_the_template_offers_the_options_we_write():
    """If Verra rewords an option, writing our string would render as free text
    in the control and read as an override."""
    from app.services.pdd_builder import _resolve

    doc = docx.Document(str(_resolve("VCS-ESG-Risk-Assessment-Template-v5.0.docx")))
    tables = find_category_tables(doc)
    row = tables[RiskCategory.E1].risk_rows[0]
    controls = _sdt_elements(row.cells[2])
    assert len(controls) >= 3

    likelihood = dropdown_options(controls[0])
    severity = dropdown_options(controls[1])
    for value in LIKELIHOOD_OPTIONS.values():
        assert value in likelihood
    for value in SEVERITY_OPTIONS.values():
        assert value in severity


# --- the dropdowns actually carry the value -------------------------------

def test_dropdowns_are_set_not_just_the_cell_text(rendered):
    """The value must live inside w:sdtContent. Text placed beside the control
    leaves Word displaying 'Select'."""
    doc, _ = rendered
    tables = find_category_tables(doc)
    controls = _sdt_elements(tables[RiskCategory.E1].risk_rows[0].cells[2])

    values = []
    for sdt in controls[:3]:
        content = sdt.find(qn("w:sdtContent"))
        node = content.find(f".//{qn('w:t')}")
        values.append(node.text if node is not None else "")

    assert values[0] == LIKELIHOOD_OPTIONS[3]
    assert values[1] == SEVERITY_OPTIONS[2]
    assert values[2] == "Low"        # severity 2 x likelihood 3 per the matrix


def test_no_dropdown_still_reads_select_where_we_wrote(rendered):
    doc, _ = rendered
    tables = find_category_tables(doc)
    for category in (RiskCategory.E1, RiskCategory.S1, RiskCategory.G2):
        cell = tables[category].risk_rows[0].cells[2]
        for sdt in _sdt_elements(cell)[:3]:
            content = sdt.find(qn("w:sdtContent"))
            node = content.find(f".//{qn('w:t')}")
            assert (node.text or "") != "Select"


def test_the_risk_level_matches_the_engine_not_the_author(rendered):
    """The level is computed from the matrix, never typed."""
    doc, _ = rendered
    tables = find_category_tables(doc)
    controls = _sdt_elements(tables[RiskCategory.S1].risk_rows[0].cells[2])
    content = controls[2].find(qn("w:sdtContent"))
    assert content.find(f".//{qn('w:t')}").text == "Low"


# --- row content -----------------------------------------------------------

def test_the_risk_id_and_mitigation_are_written(rendered):
    doc, _ = rendered
    tables = find_category_tables(doc)
    row = tables[RiskCategory.E1].risk_rows[0]
    assert row.cells[0].text.strip() == "E1.1"
    assert "fenced footprint" in row.cells[3].text


def test_the_justification_reaches_the_level_cell(rendered):
    doc, _ = rendered
    tables = find_category_tables(doc)
    cell = tables[RiskCategory.E1].risk_rows[0].cells[2]
    assert "degraded agricultural land" in cell.text


def test_a_not_applicable_category_is_marked_with_its_reason(tmp_path):
    assessment = _assessment(**{
        RiskCategory.S6: {"not_applicable": True,
                          "na_justification": "No armed personnel are engaged."}
    })
    path = tmp_path / "esg.docx"
    render_esg(assessment, path)
    doc = docx.Document(str(path))
    cell = find_category_tables(doc)[RiskCategory.S6].risk_rows[0].cells[2]
    assert "N/A" in cell.text
    assert "armed personnel" in cell.text.lower()


# --- what is left for the author ------------------------------------------

def test_verras_own_remaining_questions_are_reported(rendered):
    """Verra pre-writes 44 risk questions; we assess one per category. The rest
    are outstanding and must be visible as such."""
    _, report = rendered
    assert report.total_rows_left > 20
    assert report.rows_left_for_author["S1"] == 7   # 8 rows, 1 written


def test_a_category_with_no_table_is_reported_not_dropped(rendered):
    """G4 has no risk table in the v5.0 template. Silently dropping it would
    let someone believe an assessed category reached the document."""
    _, report = rendered
    assert any("G4" in c for c in report.categories_without_table)
    assert any("G4" in w for w in report.warnings)


def test_the_report_reads_as_a_to_do_list(rendered):
    _, report = rendered
    text = report.as_text()
    assert "still need an author" in text
    assert "S1" in text


def test_the_source_template_is_never_modified(tmp_path):
    from app.services.pdd_builder import _resolve

    source = _resolve("VCS-ESG-Risk-Assessment-Template-v5.0.docx")
    before = source.read_bytes()
    render_esg(_assessment(), tmp_path / "esg.docx")
    assert source.read_bytes() == before


def test_the_output_opens_as_a_word_document(tmp_path):
    path = tmp_path / "esg.docx"
    render_esg(_assessment(), path)
    assert path.stat().st_size > 20_000
    assert len(docx.Document(str(path)).tables) > 10


# --- rows Verra pre-marked N/A -------------------------------------------

def test_a_row_verra_marked_not_applicable_is_not_overwritten(rendered):
    """S2's first row is pre-marked N/A and carries no dropdowns. Writing an
    assessment into it would overwrite a determination Verra made and record no
    severity or likelihood at all."""
    doc, report = rendered
    table = find_category_tables(doc)[RiskCategory.S2]
    first = table.risk_rows[0]
    assert "N/A" in first.cells[2].text
    assert first.cells[0].text.strip() != "S2.1"


def test_the_assessment_lands_in_the_first_usable_row(rendered):
    doc, _ = rendered
    table = find_category_tables(doc)[RiskCategory.S2]
    target = table.assessable_rows[0]
    assert target.cells[0].text.strip() == "S2.1"
    assert len(_sdt_elements(target.cells[2])) >= 3


def test_no_dropdown_warning_remains(rendered):
    """The 'expected 3 dropdowns, found 0' warning meant we were writing into a
    pre-marked N/A row."""
    _, report = rendered
    assert not any("expected 3 dropdowns" in w for w in report.warnings)
