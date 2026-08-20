"""Module 3 tests — PDD content model and docx rendering."""

from __future__ import annotations

from datetime import date

import docx
import pytest

from app.domain import constants as K
from app.domain.additionality import FinancialInputs, assess_additionality
from app.domain.baseline import emission_reductions
from app.domain.classification import ProjectIntake, Severity, classify
from app.domain.emission_factors import PowerUnit, grid_emission_factor
from app.domain.pdd_content import (
    ProjectIdentity,
    build_pdd_content,
    fmt_date,
)
from app.services import docx_filler
from app.services.pdd_builder import build_pdd, template_path


def _intake(**overrides) -> ProjectIntake:
    base = dict(
        name="Aligarh Solar One",
        proponent="Bodhi Hub Client",
        country_iso2="IN",
        technology=K.Technology.SOLAR_PV_TERRESTRIAL,
        installed_capacity_mw=50.0,
        expected_annual_generation_mwh=87_600.0,
        initial_crediting_period_start=date(2026, 3, 1),
    )
    base.update(overrides)
    return ProjectIntake(**base)


def _grid() -> list[PowerUnit]:
    return [
        PowerUnit("COAL-1", 50_000, 2012, efficiency=0.35,
                  efficiency_fuel_ef_t_per_gj=0.0946),
        PowerUnit("COAL-2", 30_000, 2021, efficiency=0.38,
                  efficiency_fuel_ef_t_per_gj=0.0946),
        PowerUnit("GAS-1", 10_000, 2023, efficiency=0.52,
                  efficiency_fuel_ef_t_per_gj=0.0561),
        PowerUnit("HYDRO-1", 10_000, 2008, low_cost_must_run=True,
                  generation_only=True),
    ]


def _full_content(intake=None, **kw):
    intake = intake or _intake()
    cls = classify(intake)
    ef = grid_emission_factor(_grid(), intake.technology, 1)
    er = emission_reductions(intake.expected_annual_generation_mwh,
                             ef.ef_grid_cm, eg_facility_mwh=87_600,
                             ef_embodied_kg_per_mwh=25.0)
    add = assess_additionality(
        FinancialInputs(
            capex=40_000, annual_opex=500,
            annual_generation_mwh=87_600, tariff_per_mwh=0.0300,
            project_lifetime_years=25, discount_rate=0.10, benchmark_irr=0.14,
            credit_price_per_tco2e=0.0080,
            annual_credits_tco2e=er.emission_reductions_tco2e),
        n_all=10, n_diff=9, project_capacity_mw=50.0, regulatory_surplus=True)
    return build_pdd_content(intake, cls, ProjectIdentity(**kw), ef, er, add)


def _finding(findings, check):
    return next(f for f in findings if f.check == check)


# --- date formatting -------------------------------------------------------

def test_dates_use_the_verra_dd_mmm_yyyy_format():
    assert fmt_date(date(2026, 3, 1)) == "01-MAR-2026"


# --- template routing ------------------------------------------------------

def test_template_a_before_the_cutover():
    content = _full_content(_intake(initial_crediting_period_start=date(2026, 12, 31)))
    assert content.template_version is K.TemplateVersion.A


def test_template_b_on_the_cutover():
    content = _full_content(_intake(initial_crediting_period_start=date(2027, 1, 1)))
    assert content.template_version is K.TemplateVersion.B


def test_both_official_templates_are_present():
    for version in (K.TemplateVersion.A, K.TemplateVersion.B):
        assert template_path(version).exists()


# --- content model ---------------------------------------------------------

def test_cover_page_fields_are_populated():
    content = _full_content()
    assert content.fields["Project name"] == "Aligarh Solar One"
    assert content.fields["Crediting period start"] == "01-MAR-2026"
    assert content.fields["Crediting period end"] == "01-MAR-2031"
    assert content.fields["Methodology ID and version"] == "VMR0017 v1.0"


def test_annual_estimates_span_the_crediting_period_not_the_project_life():
    """Five rows, not 25: the crediting period is what generates credits."""
    content = _full_content()
    assert len(content.annual_estimates) == K.EI_CREDITING_PERIOD_YEARS


def test_total_reductions_are_the_sum_of_annual_rows():
    content = _full_content()
    assert content.total_estimated_reductions == pytest.approx(
        sum(a.reductions_tco2e for a in content.annual_estimates))


def test_flat_annual_estimates_are_flagged():
    """VT0011 para 72 Option 2 requires an annually updated build margin."""
    content = _full_content()
    assert _finding(
        content.findings, "pdd.annual_estimates").severity is Severity.WARNING


def test_quantification_prose_carries_the_actual_numbers():
    content = _full_content()
    text = " ".join(content.sections["Baseline Emissions"])
    assert "BE_y = EG_PJ,y" in text
    assert "tCO2e per year" in text


def test_additionality_prose_records_the_barrier_analysis_exclusion():
    content = _full_content()
    text = " ".join(content.sections["Additionality Methods"])
    assert "5.3.2" in text
    assert "barrier analysis" in text.lower()


def test_ccp_ineligibility_appears_in_the_prose():
    content = _full_content()
    text = " ".join(content.sections["Additionality Methods"]).lower()
    assert "core carbon principles" in text


# --- refusal to build on bad inputs ---------------------------------------

def test_missing_engine_outputs_block_the_build():
    intake = _intake()
    content = build_pdd_content(intake, classify(intake))
    assert content.blocked
    for check in ("pdd.emission_factors", "pdd.quantification",
                  "pdd.additionality"):
        assert _finding(content.findings, check).severity is Severity.FAIL


def test_blocked_classification_blocks_the_pdd():
    """A high-income country fails VMR0017 Table 1; no PDD may claim otherwise."""
    intake = _intake(country_iso2="DE")
    content = build_pdd_content(intake, classify(intake))
    assert _finding(content.findings, "pdd.classification").severity is Severity.FAIL


def test_non_additional_verdict_blocks_the_pdd():
    intake = _intake()
    cls = classify(intake)
    ef = grid_emission_factor(_grid(), intake.technology, 1)
    er = emission_reductions(87_600, ef.ef_grid_cm, eg_facility_mwh=87_600,
                             ef_embodied_kg_per_mwh=25.0)
    add = assess_additionality(
        FinancialInputs(
            capex=40_000, annual_opex=500, annual_generation_mwh=87_600,
            tariff_per_mwh=0.0300, project_lifetime_years=25,
            discount_rate=0.10, benchmark_irr=0.14,
            credit_price_per_tco2e=0.0080, annual_credits_tco2e=71_000),
        n_all=10, n_diff=1,  # common practice -> not additional
        project_capacity_mw=50.0, regulatory_surplus=True)
    content = build_pdd_content(intake, cls, ProjectIdentity(), ef, er, add)
    assert _finding(content.findings, "pdd.additionality").severity is Severity.FAIL


# --- docx rendering --------------------------------------------------------

def test_build_writes_a_readable_docx(tmp_path):
    result = build_pdd(_full_content(), tmp_path / "PDD.docx")
    assert result.output_path.exists()
    doc = docx.Document(str(result.output_path))
    assert len(doc.paragraphs) > 100


def test_every_cover_page_field_is_written(tmp_path):
    result = build_pdd(_full_content(), tmp_path / "PDD.docx")
    assert result.report.fields_not_found == []
    assert result.report.completion_ratio == 1.0


def test_every_drafted_section_is_located(tmp_path):
    result = build_pdd(_full_content(), tmp_path / "PDD.docx")
    assert result.report.sections_not_found == []


def test_project_name_appears_in_the_rendered_document(tmp_path):
    result = build_pdd(_full_content(), tmp_path / "PDD.docx")
    doc = docx.Document(str(result.output_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Aligarh Solar One" in text


def test_guidance_under_a_filled_section_is_replaced(tmp_path):
    result = build_pdd(_full_content(), tmp_path / "PDD.docx")
    doc = docx.Document(str(result.output_path))
    _, instructions = docx_filler._section_paragraphs(doc, "Baseline Scenario")
    assert instructions == []


def test_unfilled_sections_keep_their_guidance(tmp_path):
    """The remaining guidance IS the to-do list; it must survive a draft build."""
    result = build_pdd(_full_content(), tmp_path / "PDD.docx")
    assert sum(result.report.instructions_remaining.values()) > 50
    assert _finding(result.findings, "pdd.completion").severity is Severity.WARNING


def test_completion_report_ranks_the_biggest_gaps_first(tmp_path):
    result = build_pdd(_full_content(), tmp_path / "PDD.docx")
    gaps = result.sections_needing_input
    assert gaps == sorted(gaps, key=lambda kv: -kv[1])


def test_strip_guidance_produces_a_clean_submission_copy(tmp_path):
    result = build_pdd(_full_content(), tmp_path / "final.docx",
                       strip_guidance=True)
    doc = docx.Document(str(result.output_path))
    assert not [p for p in doc.paragraphs
                if p.style.name == docx_filler.INSTRUCTION_STYLE]
    assert result.report.instructions_remaining == {}


def test_the_source_template_is_never_modified(tmp_path):
    source = template_path(K.TemplateVersion.A)
    before = source.read_bytes()
    build_pdd(_full_content(), tmp_path / "PDD.docx", strip_guidance=True)
    assert source.read_bytes() == before


def test_repeated_labels_are_written_only_once(tmp_path):
    """'Justification' appears in many tables; a naive fill would overwrite
    every one of them with the same value."""
    result = build_pdd(_full_content(), tmp_path / "PDD.docx")
    doc = docx.Document(str(result.output_path))
    values = [r.cells[1].text.strip()
              for t in doc.tables for r in t.rows
              if len(r.cells) >= 2 and r.cells[0].text.strip() == "Project name"]
    assert values.count("Aligarh Solar One") == 1


# --- an absent figure must not become a favourable claim -------------------

def _additionality_prose(add):
    from app.domain.pdd_content import _additionality

    return " ".join(_additionality(add))


def test_an_unassessed_common_practice_is_reported_as_outstanding():
    """Writing "0 similar projects identified" would put a favourable claim in
    the Project Description that nobody established."""
    from app.domain.additionality import FinancialInputs, assess_additionality

    add = assess_additionality(
        FinancialInputs(capex=25_000, annual_opex=500,
                        annual_generation_mwh=87_600, tariff_per_mwh=0.03,
                        project_lifetime_years=25, discount_rate=0.10,
                        benchmark_irr=0.14, credit_price_per_tco2e=0.008,
                        annual_credits_tco2e=None),
        n_all=None, n_diff=0, project_capacity_mw=50.0,
        regulatory_surplus=True)

    prose = _additionality_prose(add)
    assert "has not yet been completed" in prose
    assert "0 similar projects" not in prose


def test_an_out_of_range_return_is_not_written_as_no_cashflow():
    """A missing IRR has two opposite causes. Out of range means the project is
    wildly profitable; no sign change means it never turns positive. Only the
    second supports additionality."""
    from app.domain.additionality import FinancialInputs, assess_additionality

    add = assess_additionality(
        # Mixed units: capex in lakh against a tariff in rupees.
        FinancialInputs(capex=40_000, annual_opex=500,
                        annual_generation_mwh=87_600, tariff_per_mwh=3_000,
                        project_lifetime_years=25, discount_rate=0.10,
                        benchmark_irr=0.14, credit_price_per_tco2e=0.008,
                        annual_credits_tco2e=None),
        n_all=None, n_diff=0, project_capacity_mw=50.0,
        regulatory_surplus=True)

    prose = _additionality_prose(add)
    assert "outside a plausible range" in prose
    assert "no positive net cashflow" not in prose


def test_the_content_builder_survives_every_absent_figure():
    """The combination that broke document-status in production: no grid data,
    no similar-project search, mixed units."""
    from datetime import date

    from app.domain.additionality import FinancialInputs, assess_additionality
    from app.domain.classification import ProjectIntake, classify
    from app.domain.constants import Technology
    from app.domain.pdd_content import ProjectIdentity, build_pdd_content

    intake = ProjectIntake(
        name="Aligarh Solar One", proponent="Bodhi Hub Client",
        country_iso2="IN", technology=Technology.SOLAR_PV_TERRESTRIAL,
        installed_capacity_mw=50.0, expected_annual_generation_mwh=87_600.0,
        initial_crediting_period_start=date(2026, 3, 1))
    add = assess_additionality(
        FinancialInputs(capex=40_000, annual_opex=500,
                        annual_generation_mwh=87_600, tariff_per_mwh=0.03,
                        project_lifetime_years=25, discount_rate=0.10,
                        benchmark_irr=0.14, credit_price_per_tco2e=0.008,
                        annual_credits_tco2e=None),
        n_all=None, n_diff=0, project_capacity_mw=50.0,
        regulatory_surplus=True)

    content = build_pdd_content(intake, classify(intake), ProjectIdentity(),
                                None, None, add)
    assert content.sections["Additionality Methods"]
